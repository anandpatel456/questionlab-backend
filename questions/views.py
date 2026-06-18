from django.db.models import Count, Sum
from django.views import View
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.db import models as django_models
from rest_framework.views import APIView
from rest_framework.response import Response
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from io import BytesIO
from django.http import FileResponse

from drf_spectacular.utils import extend_schema
from .services import generate_questions, save_question_bank
from .models import Difficulty, Question, QuestionBank, QuestionType, Subject, Concept, Activity
from .serializers import (
    GenerateQuestionsSerializer,
    QuestionBankDetailSerializer,
    QuestionBankSerializer,
    QuestionSerializer,
    SaveQuestionBankSerializer,
)
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document


@extend_schema(
    request=GenerateQuestionsSerializer,
    responses={200: dict},
)
class GenerateQuestionsView(APIView):

    def post(self, request):
        serializer = GenerateQuestionsSerializer(
            data=request.data
        )
        serializer.is_valid(raise_exception=True)
        questions = generate_questions(serializer.validated_data)
        Activity.objects.create(
            action_type="generate",
            description=f"Generated {len(questions)} questions",
            subject=serializer.validated_data.get("subject", ""),
        )
        return Response({
            "success": True,
            "questions": questions
        })


class MasterDataView(APIView):

    def get(self, request):
        subjects = list(Subject.objects.values_list("name", flat=True))
        difficulties = list(Difficulty.objects.values_list("name", flat=True))
        question_types = list(QuestionType.objects.values("value", "name"))
        return Response({
            "subjects": subjects,
            "difficulties": difficulties,
            "question_types": question_types,
        })


class QuestionBankListView(APIView):

    def get(self, request):
        banks = QuestionBank.objects.annotate(question_count=Count("question"))
        serializer = QuestionBankSerializer(banks, many=True)
        return Response(serializer.data)


class SaveQuestionBankView(APIView):

    def post(self, request):
        serializer = SaveQuestionBankSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        bank = save_question_bank(
            serializer.validated_data["subject"],
            serializer.validated_data["questions"],
            serializer.validated_data.get("concepts", []),
            serializer.validated_data.get("bank_name"),
        )
        Activity.objects.create(
            action_type="bank",
            description=f"Created question bank '{bank.name}'",
            subject=serializer.validated_data.get("subject", ""),
        )
        return Response({
            "success": True,
            "bank_id": bank.id
        })


class ConceptListView(APIView):

    def get(self, request):
        subject_name = request.GET.get("subject")
        concepts = Concept.objects.filter(subject__name=subject_name)
        return Response({
            "concepts": [concept.name for concept in concepts]
        })


class DeleteQuestionBankView(APIView):

    def delete(self, request, bank_id):
        bank = get_object_or_404(QuestionBank, id=bank_id)
        Activity.objects.create(
            action_type="delete",
            description=f"Deleted question bank '{bank.name}'",
            subject=bank.name,
        )
        bank.delete()
        return Response({"success": True})


class QuestionBankDetailView(APIView):

    def get(self, request, pk):
        bank = get_object_or_404(QuestionBank, id=pk)
        questions = Question.objects.filter(question_bank=bank)
        serializer = QuestionSerializer(questions, many=True)
        return Response({
            "id": bank.id,
            "name": bank.name,
            "questions": serializer.data
        })


@method_decorator(csrf_exempt, name='dispatch')
class ExportQuestionBankView(View):

    def get(self, request, pk):
        bank = get_object_or_404(QuestionBank, pk=pk)
        questions = bank.question_set.all()
        format_type = request.GET.get("format", "json")

        # INCREMENT EXPORT COUNT FOR ALL FORMATS
        QuestionBank.objects.filter(pk=pk).update(
            export_count=django_models.F('export_count') + 1
        )
        Activity.objects.create(
            action_type="export",
            description=f"Exported {format_type.upper()} '{bank.name}'",
            subject=bank.name,
        )

        if format_type == "pdf":
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer)
            styles = getSampleStyleSheet()
            elements = []
            elements.append(Paragraph(bank.name, styles["Title"]))
            elements.append(Spacer(1, 12))
            for index, q in enumerate(questions, start=1):
                elements.append(Paragraph(f"{index}. {q.question_text}", styles["Heading3"]))
                for option in q.options:
                    elements.append(Paragraph(option, styles["BodyText"]))
                elements.append(Paragraph(f"Answer: {q.answer}", styles["BodyText"]))
                elements.append(Spacer(1, 10))
            doc.build(elements)
            buffer.seek(0)
            return FileResponse(buffer, as_attachment=True, filename=f"{bank.name}.pdf", content_type="application/pdf")

        if format_type == "docx":
            document = Document()
            document.add_heading(bank.name, level=1)
            for index, q in enumerate(questions, start=1):
                document.add_paragraph(f"{index}. {q.question_text}")
                for option in q.options:
                    document.add_paragraph(option)
                document.add_paragraph(f"Answer: {q.answer}")
                document.add_paragraph("")
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return FileResponse(buffer, as_attachment=True, filename=f"{bank.name}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        data = {
            "id": bank.id,
            "name": bank.name,
            "questions": [
                {"question": q.question_text, "options": q.options, "answer": q.answer}
                for q in questions
            ],
        }
        return JsonResponse(data, safe=False)


class UpdateQuestionView(APIView):
    def patch(self, request, question_id):
        question = get_object_or_404(Question, id=question_id)
        if "question_text" in request.data:
            question.question_text = request.data["question_text"]
        if "options" in request.data:
            question.options = request.data["options"]
        if "answer" in request.data:
            question.answer = request.data["answer"]
        question.save()
        Activity.objects.create(
            action_type="edit",
            description=f"Edited question in '{question.question_bank.name}'",
            subject=question.question_bank.name,
        )
        return Response({
            "id": question.id,
            "question_text": question.question_text,
            "options": question.options,
            "answer": question.answer,
        })


class UpdateQuestionBankView(APIView):
    def patch(self, request, pk):
        bank = get_object_or_404(QuestionBank, pk=pk)
        if "name" in request.data:
            bank.name = request.data["name"]
            bank.save()
            Activity.objects.create(
                action_type="edit",
                description=f"Edited question bank name to '{bank.name}'",
                subject=bank.name,
            )
        return Response({"id": bank.id, "name": bank.name})
    
class DashboardStatsView(APIView):
    def get(self, request):
        total_questions = Question.objects.count()
        total_banks = QuestionBank.objects.count()
        total_subjects = Subject.objects.count()
        total_exports = QuestionBank.objects.aggregate(
            total=Sum('export_count')
        )['total'] or 0

        subjects = Subject.objects.annotate(
            question_count=Count('question')
        ).order_by('-question_count')[:5]

        max_count = subjects[0].question_count if subjects else 1

        top_subjects = [
            {
                "name": s.name,
                "count": s.question_count,
                "percentage": round((s.question_count / max_count) * 100) if max_count > 0 else 0,
            }
            for s in subjects
        ]


        recent = Activity.objects.all()[:5]
        activities = [
            {
                "action": a.description,
                "subject": a.subject,
                "time": a.created_at.isoformat(),
                "type": a.action_type,
            }
            for a in recent
        ]

        return Response({
            "total_questions": total_questions,
            "total_banks": total_banks,
            "total_subjects": total_subjects,
            "total_exports": total_exports,
            "recent_activity": activities,
            "top_subjects": top_subjects,
        })

class ActivityListView(APIView):
    def get(self, request):
        activities = Activity.objects.all()
        data = [
            {
                "action": a.description,
                "subject": a.subject,
                "time": a.created_at.isoformat(),
                "type": a.action_type,
            }
            for a in activities
        ]
        return Response({"activities": data})